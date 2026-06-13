import io
import logging
from fpdf import FPDF
from docx import Document

logger = logging.getLogger(__name__)

def markdown_to_pdf(markdown_text: str) -> bytes:
    """
    Parses a basic markdown string and generates a clean PDF document using fpdf2.
    """
    logger.info("Converting Markdown to PDF...")
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Enable margins
        pdf.set_margins(15, 15, 15)
        
        lines = markdown_text.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                pdf.ln(4)
                continue
                
            # Headers
            if stripped.startswith("# "):
                pdf.set_font("Helvetica", style="B", size=18)
                pdf.cell(0, 10, stripped[2:], ln=True)
                pdf.ln(4)
            elif stripped.startswith("## "):
                pdf.set_font("Helvetica", style="B", size=14)
                pdf.cell(0, 8, stripped[3:], ln=True)
                pdf.ln(3)
            elif stripped.startswith("### "):
                pdf.set_font("Helvetica", style="B", size=12)
                pdf.cell(0, 6, stripped[4:], ln=True)
                pdf.ln(2)
            # Bullet items
            elif stripped.startswith("- ") or stripped.startswith("* "):
                pdf.set_font("Helvetica", size=10)
                # Bullet indentation
                pdf.cell(5, 5, "*", ln=False)
                pdf.multi_cell(0, 5, stripped[2:])
                pdf.ln(1)
            # Standard paragraphs
            else:
                pdf.set_font("Helvetica", size=10)
                # Replace special characters that aren't natively supported in standard Latin-1 Helvetica
                clean_line = stripped.replace("—", "-").replace("–", "-").replace("’", "'").replace("“", '"').replace("”", '"')
                pdf.multi_cell(0, 6, clean_line)
                pdf.ln(2)
                
        # Return raw PDF bytes
        return bytes(pdf.output())
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        # Return fallback error PDF text
        fallback = FPDF()
        fallback.add_page()
        fallback.set_font("Helvetica", size=12)
        fallback.cell(0, 10, f"Error generating PDF: {str(e)}", ln=True)
        return bytes(fallback.output())

def markdown_to_docx(markdown_text: str) -> bytes:
    """
    Parses a basic markdown string and generates a Microsoft Word document using python-docx.
    """
    logger.info("Converting Markdown to DOCX...")
    try:
        doc = Document()
        lines = markdown_text.split("\n")
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
                
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)
                
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        doc = Document()
        doc.add_heading("Error generating report", level=1)
        doc.add_paragraph(str(e))
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()
